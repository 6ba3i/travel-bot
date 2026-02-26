#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Installation (pip):
  pip install -U faster-whisper python-docx openpyxl tqdm

Optional system dependency (for reading duration via ffprobe fallback):
  - Install FFmpeg (ffprobe/ffmpeg in PATH)

Usage example:
  python transcribe_rh_m4a.py --input-dir /mnt/data --model large-v3 --device auto
"""

from __future__ import annotations

import argparse
import os
import re
import sys
import traceback
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

CONTEXT_TEXT = (
    "PHF COM, société marocaine spécialisée dans l’intégration de systèmes "
    "audiovisuels, avec aussi une présence en France, créée en 2009 à Casablanca. "
    "De 4 collaborateurs à près de 20 en 2025. Signes: démotivation, nonchalance, "
    "promesses non tenues, manque de suivi et professionnalisme. Axes: motivation, "
    "discipline, compétences. Objectifs: identifier facteurs réels, comprendre freins "
    "organisationnels et managériaux, perception et projection collaborateurs, qualité "
    "communication interne et climat social, procédures de travail, plan d’action "
    "opérationnel. Acteurs clés: Younes Yamouni (DG fondateur), Ali (DGA commercial "
    "actionnaire, arrivé en 2022 selon entretien), Amina (DAF), Houssam (Directeur "
    "technique)."
)

TRANSCRIPTION_RULES = [
    "Transcription brute fidèle au contenu oral (français avec mots en darija conservés).",
    "Ne pas corriger le fond ni reformuler: conserver hésitations et tournures quand utiles.",
    "Timestamps par segment au format [mm:ss].",
    "Signaler les passages inaudibles par [inaudible] si nécessaire.",
]


@dataclass
class AudioFile:
    path: Path
    speaker: str
    part_number: Optional[int]
    display_name: str


@dataclass
class SegmentLine:
    start_s: float
    text: str


def eprint(*args, **kwargs):
    print(*args, file=sys.stderr, **kwargs)


def require_dependencies():
    missing = []
    try:
        import faster_whisper  # noqa: F401
    except Exception:
        missing.append("faster-whisper")

    try:
        import docx  # noqa: F401
    except Exception:
        missing.append("python-docx")

    try:
        import openpyxl  # noqa: F401
    except Exception:
        missing.append("openpyxl")

    if missing:
        raise RuntimeError(
            "Dépendances manquantes: "
            + ", ".join(missing)
            + ". Installez-les avec: pip install -U faster-whisper python-docx openpyxl tqdm"
        )


def sanitize_filename(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def strip_datetime_prefix(filename_stem: str) -> str:
    # Supprime un préfixe du type "mar., 10.13 " ou variantes proches
    pattern = r"^\s*[a-zA-Zéûîôàç\.]+,?\s*\d{1,2}[\.:]\d{2}\s+"
    return re.sub(pattern, "", filename_stem, flags=re.IGNORECASE).strip()


def extract_speaker_and_part(filename_stem: str) -> Tuple[str, Optional[int], str]:
    base = strip_datetime_prefix(filename_stem)
    normalized = re.sub(r"\s+", " ", base).strip()

    # Cas: "ALI NYA1" ou "ALI NYA 2" => speaker = "ALI NYA", part = 1|2
    m = re.match(r"^(.*?\bNYA)\s*([12])$", normalized, flags=re.IGNORECASE)
    if m:
        speaker = m.group(1).strip()
        part = int(m.group(2))
        display = f"{speaker}"
        return speaker.upper(), part, display

    # Cas plus générique: terminaison numérique isolée
    m2 = re.match(r"^(.*?)\s+([0-9]+)$", normalized)
    if m2:
        speaker = m2.group(1).strip()
        part = int(m2.group(2))
        return speaker.upper(), part, speaker

    return normalized.upper(), None, normalized


def format_mmss(seconds: float) -> str:
    s = max(0, int(seconds))
    mm = s // 60
    ss = s % 60
    return f"{mm:02d}:{ss:02d}"


def get_audio_duration_seconds(path: Path) -> Optional[float]:
    # Essai via mutagen (si dispo), sinon ffprobe
    try:
        from mutagen import File as MutagenFile  # type: ignore

        audio = MutagenFile(path)
        if audio is not None and hasattr(audio, "info") and hasattr(audio.info, "length"):
            return float(audio.info.length)
    except Exception:
        pass

    try:
        import json
        import subprocess

        cmd = [
            "ffprobe",
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "json",
            str(path),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, check=True)
        data = json.loads(proc.stdout)
        return float(data["format"]["duration"])
    except Exception:
        return None


def transcribe_audio(model, file_path: Path, vad_filter: bool = True) -> List[SegmentLine]:
    segments, _info = model.transcribe(
        str(file_path),
        language="fr",
        vad_filter=vad_filter,
        beam_size=5,
        best_of=5,
        condition_on_previous_text=True,
        word_timestamps=False,
        temperature=0.0,
    )

    out: List[SegmentLine] = []
    for seg in segments:
        text = (seg.text or "").strip()
        if not text:
            continue
        out.append(SegmentLine(start_s=float(seg.start), text=text))
    return out


def build_text_block(segments: List[SegmentLine]) -> str:
    lines = []
    for seg in segments:
        lines.append(f"[{format_mmss(seg.start_s)}] {seg.text}")
    return "\n".join(lines)


def write_txt(out_path: Path, title: str, speaker: str, parts_data: List[Tuple[str, str]]):
    with out_path.open("w", encoding="utf-8") as f:
        f.write(title + "\n")
        f.write("=" * len(title) + "\n\n")
        f.write("Contexte\n--------\n")
        f.write(CONTEXT_TEXT + "\n\n")
        f.write("Règles de transcription\n-----------------------\n")
        for rule in TRANSCRIPTION_RULES:
            f.write(f"- {rule}\n")
        f.write("\n")

        f.write("Transcription brute\n-------------------\n")
        for part_title, text in parts_data:
            f.write(f"\n{part_title}\n")
            f.write("~" * len(part_title) + "\n")
            f.write(text + "\n")


def write_docx(out_path: Path, title: str, parts_data: List[Tuple[str, str]]):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)

    doc.add_heading("Contexte", level=1)
    doc.add_paragraph(CONTEXT_TEXT)

    doc.add_heading("Règles de transcription", level=1)
    for rule in TRANSCRIPTION_RULES:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("Transcription brute", level=1)
    for part_title, text in parts_data:
        doc.add_heading(part_title, level=2)
        for line in text.splitlines():
            doc.add_paragraph(line)

    doc.save(out_path)


def write_index_excel(rows: List[Dict[str, str]], out_path: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Index"

    headers = ["Intervenant", "Fichier audio", "Durée (s)", "Taille (octets)", "Chemin", "Statut", "Message"]
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True)

    for row in rows:
        ws.append([
            row.get("speaker", ""),
            row.get("audio", ""),
            row.get("duration", ""),
            row.get("size", ""),
            row.get("path", ""),
            row.get("status", ""),
            row.get("message", ""),
        ])

    for col in ["A", "B", "C", "D", "E", "F", "G"]:
        ws.column_dimensions[col].width = 28 if col != "E" else 60

    wb.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Transcription des entretiens RH (.m4a) en français avec timestamps.")
    parser.add_argument("--input-dir", default="/mnt/data", help="Répertoire d'entrée contenant les .m4a")
    parser.add_argument("--output-dir", default=None, help="Répertoire de sortie (défaut: <input-dir>/transcriptions)")
    parser.add_argument("--model", default="large-v3", help="Modèle faster-whisper (ex: large-v3, medium)")
    parser.add_argument("--device", default="auto", help="auto|cpu|cuda")
    parser.add_argument("--compute-type", default="auto", help="auto|float16|int8|int8_float16")
    parser.add_argument("--no-vad", action="store_true", help="Désactiver le VAD")
    args = parser.parse_args()

    try:
        require_dependencies()
    except Exception as e:
        eprint(f"[ERREUR] {e}")
        sys.exit(2)

    input_dir = Path(args.input_dir)
    if not input_dir.exists() or not input_dir.is_dir():
        eprint(f"[ERREUR] Répertoire invalide: {input_dir}")
        sys.exit(2)

    output_dir = Path(args.output_dir) if args.output_dir else (input_dir / "transcriptions")
    output_dir.mkdir(parents=True, exist_ok=True)

    m4a_files = sorted(input_dir.glob("*.m4a"))
    if not m4a_files:
        print(f"Aucun fichier .m4a trouvé dans {input_dir}")
        return

    from faster_whisper import WhisperModel

    try:
        model = WhisperModel(args.model, device=args.device, compute_type=args.compute_type)
    except Exception as e:
        eprint(f"[ERREUR] Impossible de charger le modèle faster-whisper ({args.model}): {e}")
        sys.exit(2)

    groups: Dict[str, List[AudioFile]] = {}
    index_rows: List[Dict[str, str]] = []

    for path in m4a_files:
        speaker_key, part_num, display_name = extract_speaker_and_part(path.stem)
        af = AudioFile(path=path, speaker=speaker_key, part_number=part_num, display_name=display_name)
        groups.setdefault(speaker_key, []).append(af)

    for speaker_key, files in groups.items():
        # Tri: parties numériques d'abord, puis nom
        files.sort(key=lambda x: (x.part_number is None, x.part_number or 9999, x.path.name.lower()))

        parts_output: List[Tuple[str, str]] = []
        any_ok = False

        canonical_name = files[0].display_name

        for af in files:
            duration = get_audio_duration_seconds(af.path)
            size = af.path.stat().st_size if af.path.exists() else 0
            status = "OK"
            msg = ""

            try:
                segments = transcribe_audio(model, af.path, vad_filter=(not args.no_vad))
                text_block = build_text_block(segments)
                if not text_block.strip():
                    text_block = "[Aucun texte détecté]"

                if canonical_name.strip().upper() == "ALI NYA" and af.part_number in {1, 2}:
                    part_title = f"Partie {af.part_number}"
                elif af.part_number is not None:
                    part_title = f"Partie {af.part_number}"
                else:
                    part_title = f"Fichier: {af.path.name}"

                parts_output.append((part_title, text_block))
                any_ok = True
            except Exception as e:
                status = "ERREUR"
                msg = str(e)
                tb = traceback.format_exc(limit=1)
                eprint(f"[ERREUR] Transcription échouée pour {af.path.name}: {e}\n{tb}")

            index_rows.append(
                {
                    "speaker": canonical_name,
                    "audio": af.path.name,
                    "duration": f"{duration:.2f}" if duration is not None else "",
                    "size": str(size),
                    "path": str(af.path.resolve()),
                    "status": status,
                    "message": msg,
                }
            )

        if any_ok:
            title = f"Transcription entretien - {canonical_name} - NYA (ex PHF Maroc)"
            safe_name = sanitize_filename(canonical_name)
            txt_path = output_dir / f"{safe_name}.txt"
            docx_path = output_dir / f"{safe_name}.docx"

            write_txt(txt_path, title, canonical_name, parts_output)
            write_docx(docx_path, title, parts_output)
            print(f"[OK] Sorties créées pour {canonical_name}: {txt_path.name}, {docx_path.name}")
        else:
            eprint(f"[WARN] Aucun contenu transcrit avec succès pour {canonical_name}")

    index_path = output_dir / "Index_Entretiens.xlsx"
    write_index_excel(index_rows, index_path)
    print(f"[OK] Index Excel créé: {index_path}")
    print(f"Terminé à {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


if __name__ == "__main__":
    main()
